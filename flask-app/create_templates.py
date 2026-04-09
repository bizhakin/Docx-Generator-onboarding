"""
Script to generate the .docx template files for each contract type.
Run once: python create_templates.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "contract_templates")
os.makedirs(TEMPLATES_DIR, exist_ok=True)


def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_line(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    if bold_label:
        run = p.add_run(bold_label + " ")
        run.bold = True
    p.add_run(text)
    return p


def add_signature_block(doc, party1="SERVICE PROVIDER", party2="CLIENT"):
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    header = table.rows[0]
    header.cells[0].text = party1
    header.cells[1].text = party2
    for cell in header.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    labels = ["Signature:", "Name:", "Date:"]
    for i, label in enumerate(labels, 1):
        row = table.rows[i]
        row.cells[0].text = label + " ________________________"
        row.cells[1].text = label + " ________________________"


# ── 1. Service Agreement ──────────────────────────────────────────────────────
def create_service_agreement():
    doc = Document()
    set_heading(doc, "SERVICE AGREEMENT")

    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run("{{DATE}}")

    doc.add_paragraph()
    doc.add_paragraph(
        "This Service Agreement (\"Agreement\") is entered into as of {{DATE}}, between:"
    )

    add_line(doc, "{{PROVIDER_NAME}}", bold_label="Service Provider:")
    add_line(doc, "{{CLIENT_NAME}}", bold_label="Client:")
    add_line(doc, "{{CLIENT_ADDRESS}}", bold_label="Client Address:")

    set_heading(doc, "1. Services", level=2)
    p = doc.add_paragraph()
    p.add_run("The Service Provider agrees to perform the following services:\n")
    p.add_run("{{SERVICE_DESCRIPTION}}")

    set_heading(doc, "2. Term", level=2)
    doc.add_paragraph(
        "This Agreement shall commence on {{START_DATE}} and continue until {{END_DATE}}, "
        "unless earlier terminated in accordance with this Agreement."
    )

    set_heading(doc, "3. Compensation", level=2)
    doc.add_paragraph(
        "Client agrees to pay Service Provider ${{PAYMENT_AMOUNT}} for the services described herein. "
        "Payment terms: {{PAYMENT_TERMS}}."
    )

    set_heading(doc, "4. Confidentiality", level=2)
    doc.add_paragraph(
        "Each party agrees to keep confidential all non-public information received from the other party "
        "and to use such information solely for the purpose of fulfilling obligations under this Agreement."
    )

    set_heading(doc, "5. Termination", level=2)
    doc.add_paragraph(
        "Either party may terminate this Agreement with 30 days' written notice to the other party."
    )

    set_heading(doc, "6. Entire Agreement", level=2)
    doc.add_paragraph(
        "This Agreement constitutes the entire agreement between the parties with respect to its subject matter "
        "and supersedes all prior agreements, understandings, and negotiations."
    )

    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")
    doc.add_paragraph()
    add_signature_block(doc, "SERVICE PROVIDER", "CLIENT")

    path = os.path.join(TEMPLATES_DIR, "service_agreement.docx")
    doc.save(path)
    print(f"Created: {path}")


# ── 2. NDA ────────────────────────────────────────────────────────────────────
def create_nda():
    doc = Document()
    set_heading(doc, "NON-DISCLOSURE AGREEMENT")

    p = doc.add_paragraph()
    p.add_run("Effective Date: ").bold = True
    p.add_run("{{DATE}}")

    doc.add_paragraph(
        "This Non-Disclosure Agreement (\"Agreement\") is entered into as of {{DATE}}, between:"
    )
    add_line(doc, "{{DISCLOSING_PARTY}}", bold_label="Disclosing Party:")
    add_line(doc, "{{RECEIVING_PARTY}}", bold_label="Receiving Party:")

    set_heading(doc, "1. Purpose", level=2)
    doc.add_paragraph(
        "The parties wish to explore a business relationship for the following purpose:\n{{PURPOSE}}"
    )

    set_heading(doc, "2. Confidential Information", level=2)
    doc.add_paragraph(
        "\"Confidential Information\" means any data or information that is proprietary to the Disclosing Party "
        "and not generally known to the public, whether in tangible or intangible form, including but not limited to "
        "trade secrets, business plans, financial information, technical data, and customer lists."
    )

    set_heading(doc, "3. Obligations of Receiving Party", level=2)
    doc.add_paragraph(
        "The Receiving Party agrees to: (a) hold Confidential Information in strict confidence; "
        "(b) not disclose Confidential Information to any third parties without prior written consent; "
        "(c) use Confidential Information solely for the Purpose stated herein."
    )

    set_heading(doc, "4. Term", level=2)
    doc.add_paragraph(
        "This Agreement shall remain in effect for {{CONFIDENTIALITY_PERIOD}} year(s) from the Effective Date, "
        "unless earlier terminated by mutual written agreement of the parties."
    )

    set_heading(doc, "5. Governing Law", level=2)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with the laws of {{GOVERNING_LAW}}, "
        "without regard to its conflict of laws principles."
    )

    set_heading(doc, "6. Remedies", level=2)
    doc.add_paragraph(
        "The Receiving Party acknowledges that any breach of this Agreement may cause irreparable harm to the "
        "Disclosing Party for which monetary damages would be inadequate, and the Disclosing Party shall be "
        "entitled to seek equitable relief including injunction."
    )

    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()
    add_signature_block(doc, "DISCLOSING PARTY", "RECEIVING PARTY")

    path = os.path.join(TEMPLATES_DIR, "nda.docx")
    doc.save(path)
    print(f"Created: {path}")


# ── 3. Employment Offer Letter ────────────────────────────────────────────────
def create_employment_offer():
    doc = Document()
    set_heading(doc, "EMPLOYMENT OFFER LETTER")

    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run("{{DATE}}")

    doc.add_paragraph()
    add_line(doc, "{{CANDIDATE_NAME}}")
    doc.add_paragraph()

    doc.add_paragraph(
        "Dear {{CANDIDATE_NAME}},"
    )
    doc.add_paragraph()

    doc.add_paragraph(
        "On behalf of {{COMPANY_NAME}}, I am pleased to extend this offer of employment for the position of "
        "{{POSITION_TITLE}} within the {{DEPARTMENT}} department. We believe your skills and experience will "
        "be a valuable addition to our team."
    )

    set_heading(doc, "Position Details", level=2)
    add_line(doc, "{{POSITION_TITLE}}", bold_label="Title:")
    add_line(doc, "{{DEPARTMENT}}", bold_label="Department:")
    add_line(doc, "{{START_DATE}}", bold_label="Start Date:")
    add_line(doc, "{{MANAGER_NAME}}", bold_label="Reporting To:")

    set_heading(doc, "Compensation", level=2)
    add_line(doc, "${{SALARY}} per year", bold_label="Annual Salary:")
    doc.add_paragraph(
        "You will also be eligible for benefits including health insurance, paid time off, and other "
        "programs as described in the Employee Handbook."
    )

    set_heading(doc, "Conditions of Employment", level=2)
    doc.add_paragraph(
        "This offer is contingent upon satisfactory completion of a background check and your signing of "
        "the Company's standard confidentiality and intellectual property assignment agreement. Employment "
        "with {{COMPANY_NAME}} is at-will."
    )

    set_heading(doc, "Offer Expiry", level=2)
    doc.add_paragraph(
        "Please sign and return this letter by {{OFFER_EXPIRY}} to indicate your acceptance. "
        "If we do not receive your signed acceptance by that date, this offer will be withdrawn."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "We are excited about the possibility of you joining us and look forward to your positive response."
    )
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    add_line(doc, "{{MANAGER_NAME}}", bold_label="")
    add_line(doc, "{{COMPANY_NAME}}", bold_label="")

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Candidate Acceptance:").bold = True
    doc.add_paragraph("Signature: _____________________________    Date: ______________")
    add_line(doc, "{{CANDIDATE_NAME}}", bold_label="Printed Name:")

    path = os.path.join(TEMPLATES_DIR, "employment_offer.docx")
    doc.save(path)
    print(f"Created: {path}")


# ── 4. Freelance Contract ─────────────────────────────────────────────────────
def create_freelance_contract():
    doc = Document()
    set_heading(doc, "FREELANCE CONTRACT")

    p = doc.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run("{{DATE}}")

    doc.add_paragraph(
        "This Freelance Contract (\"Agreement\") is entered into as of {{DATE}}, between:"
    )
    add_line(doc, "{{CLIENT_NAME}}", bold_label="Client:")
    add_line(doc, "{{FREELANCER_NAME}}", bold_label="Freelancer:")

    set_heading(doc, "1. Project", level=2)
    add_line(doc, "{{PROJECT_NAME}}", bold_label="Project Name:")

    set_heading(doc, "2. Scope of Work", level=2)
    doc.add_paragraph("{{PROJECT_SCOPE}}")

    set_heading(doc, "3. Compensation", level=2)
    add_line(doc, "${{HOURLY_RATE}} per hour", bold_label="Hourly Rate:")
    add_line(doc, "{{ESTIMATED_HOURS}} hours", bold_label="Estimated Hours:")
    doc.add_paragraph(
        "Actual hours will be tracked and invoiced accordingly. The Freelancer shall submit invoices "
        "detailing hours worked, and Client shall remit payment per the agreed schedule: {{PAYMENT_SCHEDULE}}."
    )

    set_heading(doc, "4. Timeline", level=2)
    add_line(doc, "{{DEADLINE}}", bold_label="Project Deadline:")
    doc.add_paragraph(
        "The Freelancer shall use reasonable efforts to complete the project by the deadline. "
        "Any delays caused by Client's failure to provide required materials or approvals shall extend the timeline accordingly."
    )

    set_heading(doc, "5. Intellectual Property", level=2)
    doc.add_paragraph(
        "Upon receipt of full payment, Freelancer assigns to Client all intellectual property rights "
        "in the deliverables created under this Agreement. Freelancer retains the right to display the work "
        "in a portfolio."
    )

    set_heading(doc, "6. Independent Contractor", level=2)
    doc.add_paragraph(
        "Freelancer is an independent contractor, not an employee. Freelancer is responsible for all taxes "
        "on compensation received and shall not be entitled to employee benefits."
    )

    set_heading(doc, "7. Termination", level=2)
    doc.add_paragraph(
        "Either party may terminate this Agreement with 14 days' written notice. Client shall pay for all "
        "work completed up to the date of termination."
    )

    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")
    doc.add_paragraph()
    add_signature_block(doc, "FREELANCER", "CLIENT")

    path = os.path.join(TEMPLATES_DIR, "freelance_contract.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_service_agreement()
    create_nda()
    create_employment_offer()
    create_freelance_contract()
    print("\nAll templates created successfully!")
