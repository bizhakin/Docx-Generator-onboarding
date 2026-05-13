import os
import io
import re
import json
import uuid
import smtplib
import threading
from functools import wraps
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from datetime import date, datetime
from flask import Flask, render_template, request, jsonify, send_file, abort, session, redirect

from docx import Document
from portal_config import PORTAL_DATA

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, "contract_templates")
CONTRACTS_CONFIG = os.path.join(BASE_DIR, "contracts.json")
GENERATED_DIR = os.path.join(BASE_DIR, "generated")
os.makedirs(GENERATED_DIR, exist_ok=True)

PROVIDER_EMAIL = "Mxstermxndsbeats@gmail.com"
SMTP_USER = "Mxstermxndsbeats@gmail.com"

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "dev-fallback-key-change-in-prod")

STAFF_PASSWORD = os.environ.get("STAFF_PIN", "StudioAccess2024")

STATUS_INFO = {
    "awaiting_files":      {"label": "Awaiting Files"},
    "files_received":      {"label": "Files Received"},
    "in_production":       {"label": "In Production"},
    "draft_delivered":     {"label": "Draft Delivered"},
    "revisions":           {"label": "Revisions In Progress"},
    "revision_delivered":  {"label": "Revision Delivered"},
    "final_delivered":     {"label": "Final Delivery Sent"},
    "completed":           {"label": "Project Complete"},
}


# ── Helpers ──────────────────────────────────────────────────────────────────

def load_contracts():
    with open(CONTRACTS_CONFIG, "r") as f:
        return json.load(f)["contracts"]


def load_project_data(token):
    path = os.path.join(GENERATED_DIR, token, "project.json")
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {"status": "awaiting_files", "messages": []}


def save_project_data(token, data):
    path = os.path.join(GENERATED_DIR, token, "project.json")
    with open(path, "w") as f:
        json.dump(data, f)


def staff_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("staff_logged_in"):
            return redirect("/staff/login")
        return f(*args, **kwargs)
    return decorated


def get_all_projects():
    projects = []
    if not os.path.exists(GENERATED_DIR):
        return projects
    for token in os.listdir(GENERATED_DIR):
        meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
        if not os.path.exists(meta_path):
            continue
        with open(meta_path) as f:
            meta = json.load(f)
        project_data = load_project_data(token)
        form_data = meta.get("form_data", {})
        client_name = (
            form_data.get("Full_Name") or form_data.get("CLIENT_NAME") or
            form_data.get("Artist_Name") or "Unknown Client"
        )
        artist_name = form_data.get("Artist_Name") or form_data.get("STAGE_Name") or ""
        messages = project_data.get("messages", [])
        last_activity = (
            messages[-1]["timestamp"] if messages else
            meta.get("signed_at") or meta.get("created_at", "")
        )
        projects.append({
            "token": token,
            "meta": meta,
            "client_name": client_name,
            "artist_name": artist_name,
            "status": project_data.get("status", "awaiting_files"),
            "status_info": STATUS_INFO.get(project_data.get("status", "awaiting_files"), STATUS_INFO["awaiting_files"]),
            "messages": messages,
            "last_activity": last_activity,
            "signed": meta.get("signed", False),
        })
    projects.sort(key=lambda x: x["last_activity"], reverse=True)
    return projects


def replace_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        placeholder = "{{" + key + "}}"
        if placeholder not in paragraph.text:
            continue
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))
        if placeholder in paragraph.text:
            full_text = "".join(r.text for r in paragraph.runs)
            if placeholder in full_text:
                new_text = full_text.replace(placeholder, str(value))
                for run in paragraph.runs:
                    run.text = ""
                if paragraph.runs:
                    paragraph.runs[0].text = new_text


def clear_remaining_placeholders(paragraph):
    if "{{" not in paragraph.text:
        return
    for run in paragraph.runs:
        run.text = re.sub(r"\{\{[^}]+\}\}", "", run.text)
    if "{{" in paragraph.text:
        full_text = "".join(r.text for r in paragraph.runs)
        cleaned = re.sub(r"\{\{[^}]+\}\}", "", full_text)
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = cleaned


def fill_template(template_path, field_values):
    doc = Document(template_path)
    today = date.today().strftime("%B %d, %Y")
    field_values = dict(field_values)
    field_values["DATE"] = today
    field_values["Date"] = today

    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    for para in all_paragraphs:
        replace_in_paragraph(para, field_values)
        clear_remaining_placeholders(para)

    return doc


def get_base_url():
    domains = os.environ.get("REPLIT_DOMAINS", "")
    if domains:
        domain = domains.split(",")[0].strip()
        return f"https://{domain}"
    replit_dev = os.environ.get("REPLIT_DEV_DOMAIN", "")
    if replit_dev:
        return f"https://{replit_dev}"
    return request.host_url.rstrip("/")


def send_notification_email(token, meta, signed_path, base_url):
    """Send email to provider when client signs. Runs in background thread."""
    smtp_password = os.environ.get("SMTP_PASSWORD", "")
    if not smtp_password:
        return  # Email not configured — skip silently

    try:
        client_name = (
            meta["form_data"].get("Full_Name")
            or meta["form_data"].get("CLIENT_NAME")
            or meta["form_data"].get("CANDIDATE_NAME")
            or "Your client"
        )
        contract_name = meta.get("contract_name", "Contract")
        status_url = f"{base_url}/status/{token}"

        msg = MIMEMultipart("mixed")
        msg["From"] = SMTP_USER
        msg["To"] = PROVIDER_EMAIL
        msg["Subject"] = f"✍ {client_name} has signed — {contract_name}"

        body_html = f"""
        <div style="font-family:Arial,sans-serif;background:#000;color:#fff;padding:32px;max-width:600px;margin:auto;">
          <div style="background:#E6FB04;padding:12px 20px;display:inline-block;margin-bottom:24px;">
            <span style="font-weight:900;font-size:14px;letter-spacing:2px;color:#000;text-transform:uppercase;">
              MXSTER MXNDS PRODUCTIONS
            </span>
          </div>
          <h2 style="font-size:22px;font-weight:900;text-transform:uppercase;letter-spacing:1px;margin-bottom:8px;">
            Contract Signed
          </h2>
          <p style="color:#888;margin-bottom:24px;font-size:14px;">
            {client_name} has reviewed and signed the <strong style="color:#fff;">{contract_name}</strong>.
          </p>

          <table style="width:100%;border-collapse:collapse;margin-bottom:28px;">
            <tr>
              <td style="padding:10px 0;border-bottom:1px solid #2a2a2a;color:#888;font-size:13px;width:40%;">Client</td>
              <td style="padding:10px 0;border-bottom:1px solid #2a2a2a;color:#fff;font-size:13px;">{client_name}</td>
            </tr>
            <tr>
              <td style="padding:10px 0;border-bottom:1px solid #2a2a2a;color:#888;font-size:13px;">Contract</td>
              <td style="padding:10px 0;border-bottom:1px solid #2a2a2a;color:#fff;font-size:13px;">{contract_name}</td>
            </tr>
            <tr>
              <td style="padding:10px 0;border-bottom:1px solid #2a2a2a;color:#888;font-size:13px;">Signed On</td>
              <td style="padding:10px 0;border-bottom:1px solid #2a2a2a;color:#fff;font-size:13px;">{date.today().strftime("%B %d, %Y")}</td>
            </tr>
          </table>

          <p style="color:#888;font-size:13px;margin-bottom:16px;">
            The signed contract is attached to this email. You can also download it anytime from your status page.
          </p>

          <a href="{status_url}"
             style="display:inline-block;background:#E6FB04;color:#000;text-decoration:none;
                    padding:14px 28px;font-weight:900;font-size:12px;letter-spacing:2px;
                    text-transform:uppercase;margin-bottom:28px;">
            VIEW STATUS &amp; DOWNLOAD →
          </a>

          <p style="color:#444;font-size:11px;border-top:1px solid #2a2a2a;padding-top:16px;">
            Mxster Mxnds Productions LTD · mxstermxndsprods.store · +1 (601) 651-7869
          </p>
        </div>
        """

        msg.attach(MIMEText(body_html, "html"))

        # Attach the signed document
        with open(signed_path, "rb") as f:
            part = MIMEBase("application",
                            "vnd.openxmlformats-officedocument.wordprocessingml.document")
            part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition",
                            f'attachment; filename="signed_{meta["contract_id"]}_{date.today().isoformat()}.docx"')
            msg.attach(part)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(SMTP_USER, smtp_password)
            server.sendmail(SMTP_USER, PROVIDER_EMAIL, msg.as_string())

    except Exception as e:
        print(f"Email notification failed: {e}")


# ── Routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    contracts = load_contracts()
    return render_template("index.html", contracts=contracts)


@app.route("/fields/<contract_id>")
def get_contract_fields(contract_id):
    contracts = load_contracts()
    contract = next((c for c in contracts if c["id"] == contract_id), None)
    if not contract:
        abort(404)
    return jsonify(contract)


@app.route("/generate", methods=["POST"])
def generate():
    data = request.form.to_dict()
    contract_id = data.pop("contract_id", None)
    if not contract_id:
        abort(400)

    contracts = load_contracts()
    contract = next((c for c in contracts if c["id"] == contract_id), None)
    if not contract:
        abort(404)

    template_path = os.path.join(TEMPLATES_DIR, contract["template"])
    if not os.path.exists(template_path):
        abort(500)

    deliverables_count = int(data.pop("deliverables_count", 5))
    for i in range(1, 6):
        for field in ["Service", "Quantity", "Turnaround", "Revisions"]:
            key = f"{field}{i}"
            if i > deliverables_count:
                data[key] = ""

    token = str(uuid.uuid4())
    job_dir = os.path.join(GENERATED_DIR, token)
    os.makedirs(job_dir, exist_ok=True)

    client_sig_fields = [f["id"] for f in contract.get("client_signature_fields", [])]

    # Provider copy — client signature fields left blank
    provider_data = dict(data)
    for fid in client_sig_fields:
        provider_data[fid] = ""

    provider_doc = fill_template(template_path, provider_data)
    provider_doc.save(os.path.join(job_dir, "provider_copy.docx"))

    meta = {
        "contract_id": contract_id,
        "contract_name": contract["name"],
        "form_data": data,
        "client_signature_fields": contract.get("client_signature_fields", []),
        "template": contract["template"],
        "created_at": datetime.utcnow().isoformat(),
        "signed": False,
        "signed_at": None,
    }
    with open(os.path.join(job_dir, "meta.json"), "w") as f:
        json.dump(meta, f)

    base = get_base_url()
    return jsonify({
        "success": True,
        "token": token,
        "download_url": f"/download/{token}",
        "sign_url": f"{base}/sign/{token}",
        "status_url": f"{base}/status/{token}",
        "contract_name": contract["name"],
    })


@app.route("/download/<token>")
def download(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    path = os.path.join(GENERATED_DIR, token, "provider_copy.docx")
    if not os.path.exists(path):
        abort(404)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    contract_id = "contract"
    if os.path.exists(meta_path):
        with open(meta_path) as f:
            contract_id = json.load(f).get("contract_id", "contract")
    return send_file(path, as_attachment=True,
                     download_name=f"{contract_id}_{date.today().isoformat()}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/download-signed/<token>")
def download_signed(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    path = os.path.join(GENERATED_DIR, token, "signed_copy.docx")
    if not os.path.exists(path):
        abort(404)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    contract_id = "contract"
    if os.path.exists(meta_path):
        with open(meta_path) as f:
            contract_id = json.load(f).get("contract_id", "contract")
    return send_file(path, as_attachment=True,
                     download_name=f"{contract_id}_signed_{date.today().isoformat()}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/status/<token>")
def status_page(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)
    with open(meta_path) as f:
        meta = json.load(f)
    return render_template("status.html", token=token, meta=meta)


@app.route("/sign/<token>")
def sign_page(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)
    with open(meta_path) as f:
        meta = json.load(f)
    already_signed = meta.get("signed", False)
    return render_template("sign.html", token=token, meta=meta, already_signed=already_signed)


@app.route("/sign/<token>/submit", methods=["POST"])
def sign_submit(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    job_dir = os.path.join(GENERATED_DIR, token)
    meta_path = os.path.join(job_dir, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)
    with open(meta_path) as f:
        meta = json.load(f)

    template_path = os.path.join(TEMPLATES_DIR, meta["template"])
    if not os.path.exists(template_path):
        abort(500)

    all_data = dict(meta["form_data"])
    for sig_field in meta.get("client_signature_fields", []):
        fid = sig_field["id"]
        all_data[fid] = request.form.get(fid, "")

    doc = fill_template(template_path, all_data)
    signed_path = os.path.join(job_dir, "signed_copy.docx")
    doc.save(signed_path)

    # Update meta
    meta["signed"] = True
    meta["signed_at"] = datetime.utcnow().isoformat()
    meta["signed_by"] = {
        fid["id"]: request.form.get(fid["id"], "")
        for fid in meta.get("client_signature_fields", [])
    }
    with open(meta_path, "w") as f:
        json.dump(meta, f)

    # Send notification email in background
    base = get_base_url()
    threading.Thread(
        target=send_notification_email,
        args=(token, meta, signed_path, base),
        daemon=True
    ).start()

    return jsonify({
        "success": True,
        "portal_url": f"/portal/{token}",
        "download_url": f"/download-signed/{token}",
    })


@app.route("/portal/<token>")
def client_portal(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)
    with open(meta_path) as f:
        meta = json.load(f)

    contract_id = meta.get("contract_id", "")
    portal = dict(PORTAL_DATA.get(contract_id, {
        "scope": "Custom service agreement — see your signed contract for full details.",
        "turnaround": "To be communicated",
        "deliverables": [],
        "revisions": "Per signed agreement",
        "credits": [],
        "onboarding": [],
    }))

    # For custom_artist_package, build deliverables dynamically from form data
    if contract_id == "custom_artist_package":
        form_data = meta.get("form_data", {})
        dynamic_deliverables = []
        for i in range(1, 6):
            service = (form_data.get(f"Service{i}") or "").strip()
            if service:
                qty_parts = []
                qty = (form_data.get(f"Quantity{i}") or "").strip()
                ta = (form_data.get(f"Turnaround{i}") or "").strip()
                rev = (form_data.get(f"Revisions{i}") or "").strip()
                if qty:
                    qty_parts.append(f"Qty: {qty}")
                if ta:
                    qty_parts.append(ta)
                if rev:
                    qty_parts.append(f"{rev} revisions")
                dynamic_deliverables.append({
                    "label": service,
                    "qty": " · ".join(qty_parts) if qty_parts else "See agreement",
                })
        portal["deliverables"] = dynamic_deliverables

    form_data = meta.get("form_data", {})
    client_name = (
        form_data.get("Full_Name") or
        form_data.get("CLIENT_NAME") or
        form_data.get("Artist_Name") or
        "Client"
    )
    artist_name = form_data.get("Artist_Name") or form_data.get("STAGE_Name") or ""
    if artist_name == client_name:
        artist_name = ""
    start_date = form_data.get("START_Date") or form_data.get("Start_Date") or ""
    end_date = form_data.get("END_Date") or form_data.get("End_Date") or ""
    price = form_data.get("Price") or ""

    signed_at = meta.get("signed_at", "")
    signed_at_display = "—"
    if signed_at:
        try:
            signed_dt = datetime.fromisoformat(signed_at)
            signed_at_display = signed_dt.strftime("%B %d, %Y")
        except Exception:
            signed_at_display = signed_at[:10]

    project_data = load_project_data(token)

    return render_template(
        "portal.html",
        token=token,
        meta=meta,
        portal=portal,
        project_data=project_data,
        client_name=client_name,
        artist_name=artist_name,
        start_date=start_date,
        end_date=end_date,
        price=price,
        signed_at_display=signed_at_display,
    )


@app.route("/portal/<token>/comment", methods=["POST"])
def client_comment(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)
    with open(meta_path) as f:
        meta = json.load(f)
    text = request.form.get("text", "").strip()
    if not text:
        return redirect(f"/portal/{token}#messages")
    form_data = meta.get("form_data", {})
    client_name = (
        form_data.get("Full_Name") or form_data.get("CLIENT_NAME") or
        form_data.get("Artist_Name") or "Client"
    )
    project_data = load_project_data(token)
    project_data["messages"].append({
        "id": str(uuid.uuid4()),
        "from": "client",
        "author": client_name,
        "text": text,
        "timestamp": datetime.utcnow().isoformat(),
        "type": "message",
    })
    save_project_data(token, project_data)
    return redirect(f"/portal/{token}#messages")


@app.route("/staff/login", methods=["GET", "POST"])
def staff_login():
    error = None
    if request.method == "POST":
        if request.form.get("password") == STAFF_PASSWORD:
            session["staff_logged_in"] = True
            return redirect("/staff")
        error = "Incorrect password. Try again."
    return render_template("staff_login.html", error=error)


@app.route("/staff/logout")
def staff_logout():
    session.pop("staff_logged_in", None)
    return redirect("/staff/login")


@app.route("/staff")
@staff_required
def staff_dashboard():
    projects = get_all_projects()
    return render_template("staff.html", projects=projects, status_info=STATUS_INFO)


@app.route("/staff/<token>")
@staff_required
def staff_project(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)
    with open(meta_path) as f:
        meta = json.load(f)
    project_data = load_project_data(token)
    form_data = meta.get("form_data", {})
    client_name = (
        form_data.get("Full_Name") or form_data.get("CLIENT_NAME") or
        form_data.get("Artist_Name") or "Client"
    )
    signed_at = meta.get("signed_at", "")
    signed_at_display = "Not yet signed"
    if signed_at:
        try:
            signed_dt = datetime.fromisoformat(signed_at)
            signed_at_display = signed_dt.strftime("%B %d, %Y at %I:%M %p UTC")
        except Exception:
            signed_at_display = signed_at[:10]
    return render_template(
        "staff_project.html",
        token=token,
        meta=meta,
        project_data=project_data,
        client_name=client_name,
        signed_at_display=signed_at_display,
        status_info=STATUS_INFO,
        current_status_info=STATUS_INFO.get(project_data.get("status", "awaiting_files"), STATUS_INFO["awaiting_files"]),
    )


@app.route("/staff/<token>/status", methods=["POST"])
@staff_required
def staff_update_status(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    if not os.path.exists(os.path.join(GENERATED_DIR, token, "meta.json")):
        abort(404)
    new_status = request.form.get("status", "")
    if new_status not in STATUS_INFO:
        abort(400)
    project_data = load_project_data(token)
    project_data["status"] = new_status
    project_data["messages"].append({
        "id": str(uuid.uuid4()),
        "from": "staff",
        "author": "Studio",
        "text": f"Status updated to: {STATUS_INFO[new_status]['label']}",
        "timestamp": datetime.utcnow().isoformat(),
        "type": "status_update",
    })
    save_project_data(token, project_data)
    return redirect(f"/staff/{token}")


@app.route("/staff/<token>/message", methods=["POST"])
@staff_required
def staff_post_message(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    if not os.path.exists(os.path.join(GENERATED_DIR, token, "meta.json")):
        abort(404)
    text = request.form.get("text", "").strip()
    if not text:
        return redirect(f"/staff/{token}")
    msg_type = request.form.get("msg_type", "message")
    if msg_type not in ("message", "revision_request", "file_request",
                        "draft_delivered", "revision_delivered", "final_delivery"):
        msg_type = "message"
    author = request.form.get("author", "Studio Team").strip() or "Studio Team"
    project_data = load_project_data(token)
    project_data["messages"].append({
        "id": str(uuid.uuid4()),
        "from": "staff",
        "author": author,
        "text": text,
        "timestamp": datetime.utcnow().isoformat(),
        "type": msg_type,
    })
    auto_status = {
        "draft_delivered": "draft_delivered",
        "revision_delivered": "revision_delivered",
        "final_delivery": "final_delivered",
        "revision_request": "revisions",
    }
    if msg_type in auto_status:
        project_data["status"] = auto_status[msg_type]
    save_project_data(token, project_data)
    return redirect(f"/staff/{token}")


@app.route("/api/status/<token>")
def api_status(token):
    token = re.sub(r"[^a-f0-9\-]", "", token)
    meta_path = os.path.join(GENERATED_DIR, token, "meta.json")
    if not os.path.exists(meta_path):
        abort(404)
    with open(meta_path) as f:
        meta = json.load(f)
    return jsonify({
        "signed": meta.get("signed", False),
        "signed_at": meta.get("signed_at"),
        "contract_name": meta.get("contract_name"),
    })


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
